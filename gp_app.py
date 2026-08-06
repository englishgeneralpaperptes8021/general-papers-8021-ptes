import streamlit as st
import os
import fitz  # PyMuPDF
import gdown
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO

# --- SECRETS & CONFIGURATION ---
try:
    ADMIN_PASSWORD = st.secrets["ADMIN_PASSWORD"]
except Exception:
    ADMIN_PASSWORD = "8021Admin"

# Updated with your PYPMaterials8021 Google Drive folder ID
GD_FOLDER_ID = "11XF_9ZBu95qMcVENnEWy5P9a-BqbloA6"

# Directory names matching your Google Drive structure (Added Insert folder)
FOLDERS = {
    "June QP": "8021_June_qp",
    "Nov QP": "8021_Nov_qp",
    "June MS": "8021_June_ms",
    "Nov MS": "8021_Nov_ms",
    "Inserts": "8021_NovJune_in"
}

QP_FOLDERS = [FOLDERS["June QP"], FOLDERS["Nov QP"]]
MS_FOLDERS = [FOLDERS["June MS"], FOLDERS["Nov MS"]]
IN_FOLDERS = [FOLDERS["Inserts"]]

# Ensure local directories exist for synced files
for folder in FOLDERS.values():
    if not os.path.exists(folder):
        os.makedirs(folder)

# --- APP SETUP & STYLING ---
st.set_page_config(page_title="8021 General Paper Handout Builder", layout="wide")

custom_css = """
<style>
    /* Main Background */
    .stApp {
        background-color: #FCBBE6 !important;
    }
    
    /* Sidebar Background */
    [data-testid="stSidebar"] {
        background-color: #FA8FD6 !important;
    }

    /* Text Formatting Rules */
    html, body, [class*="css"], h1, h2, h3, h4, h5, h6, p, label, span, div {
        color: #057047 !important;
    }

    /* Input Controls */
    div[data-baseweb="input"] > div, 
    div[data-baseweb="select"] > div {
        background-color: #FEE7F6 !important;
        border-color: #F527B0 !important;
    }
    
    input {
        color: #057047 !important;
    }

    /* Buttons */
    div.stButton > button, div.stDownloadButton > button {
        background-color: #09C87F !important;
        color: #057047 !important;
        font-weight: bold !important;
        border: 1px solid #057047 !important;
        border-radius: 8px !important;
        transition: all 0.3s ease;
    }

    div.stButton > button:hover, div.stDownloadButton > button:hover {
        background-color: #63F8BF !important;
        color: #057047 !important;
        border-color: #057047 !important;
        transform: scale(1.02);
    }
    
    /* Expanders / Cards */
    div[data-testid="stExpander"] {
        background-color: rgba(255, 255, 255, 0.4) !important;
        border: 1px solid #057047 !important;
        border-radius: 8px !important;
        margin-bottom: 10px;
    }
</style>
"""
st.markdown(custom_css, unsafe_allow_html=True)

# --- HELPER FUNCTIONS ---
def sync_from_drive():
    """Downloads files from Google Drive into the local Streamlit environment."""
    try:
        with st.spinner("🔄 Syncing with Google Drive..."):
            gdown.download_folder(id=GD_FOLDER_ID, output=".", quiet=True)
        st.success("✅ Library Updated from Google Drive!")
    except Exception as e:
        st.error(f"Sync Error: {e}")

def get_filename_pattern(month, year, paper_type, paper_code):
    """Formats Cambridge PDF file patterns for 8021 (e.g., 8021_w25_in_21)."""
    short_year = year[-2:]
    month_code = 's' if month == "June" else 'w'
    return f"8021_{month_code}{short_year}_{paper_type}_{paper_code}"

def search_pdfs(keyword_list, target_folders):
    """Searches PDF contents across target folders for matching keywords."""
    results = []
    for folder_path in target_folders:
        if not os.path.exists(folder_path): 
            continue
        for file in os.listdir(folder_path):
            if file.endswith(".pdf"):
                try:
                    doc = fitz.open(os.path.join(folder_path, file))
                    for page_num in range(len(doc)):
                        text = doc[page_num].get_text().lower()
                        if all(k.lower() in text for k in keyword_list):
                            results.append({
                                "file": file, 
                                "page": page_num, 
                                "path": os.path.join(folder_path, file)
                            })
                    doc.close()
                except Exception:
                    continue
    return results

def render_page_image(pdf_path, page_num):
    """Renders a PDF page to high-res PNG image bytes."""
    doc = fitz.open(pdf_path)
    page = doc.load_page(page_num)
    pix = page.get_pixmap(matrix=fitz.Matrix(2, 2))
    img_bytes = pix.tobytes("png")
    doc.close()
    return img_bytes

def add_page_number_to_header(section):
    """Adds native dynamic Word page numbers to document header."""
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    
    fldChar1 = parse_xml(r'<w:fldChar %s w:fldCharType="begin"/>' % nsdecls('w'))
    instrText = parse_xml(r'<w:instrText %s xml:space="preserve"> PAGE </w:instrText>' % nsdecls('w'))
    fldChar2 = parse_xml(r'<w:fldChar %s w:fldCharType="separate"/>' % nsdecls('w'))
    fldChar3 = parse_xml(r'<w:fldChar %s w:fldCharType="end"/>' % nsdecls('w'))
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# --- SESSION STATE INITIALIZATION ---
if 'handout_basket' not in st.session_state:
    st.session_state.handout_basket = []
if 'search_results_qp' not in st.session_state:
    st.session_state.search_results_qp = []
if 'search_results_ms' not in st.session_state:
    st.session_state.search_results_ms = []

# --- SIDEBAR & HEADER ---
st.title("PTE SENGKURONG")
st.title("📚 8021 PYP General Paper Hub")

with st.sidebar:
    st.header("Cloud Controls")
    if st.button("🔄 Sync New Files"):
        sync_from_drive()
    st.info("Upload new PDFs into Google Drive, then click Sync here.")
    
    st.divider()
    
    basket_count = len(st.session_state.handout_basket)
    st.markdown("### 🛒 Basket Summary")
    st.metric(label="Total Pages Selected", value=f"{basket_count} page(s)")

# --- APP TABS ---
tab1, tab2, tab3, tab4, tab5 = st.tabs([
    "🔍 Search Questions", 
    "🎯 Search Answers", 
    "📅 View&Save PYP", 
    "📝 Handout Download", 
    "⚙️ Admin"
])

# --- TAB 1: SEARCH QUESTION PAPERS ---
with tab1:
    st.header("Search Question Papers only")
    col_input, col_reset = st.columns([4, 1])
    with col_input:
        keywords_qp = st.text_input("Enter keywords (e.g., 'technology')", key="input_qp")
    with col_reset:
        st.write(" ")
        if st.button("🔄 Clear Search", key="btn_clear_qp"):
            st.session_state.search_results_qp = []
            st.rerun()

    if st.button("Search Question Papers", type="primary", key="btn_search_qp"):
        if keywords_qp:
            with st.spinner("Scanning Question Papers..."):
                st.session_state.search_results_qp = search_pdfs([keywords_qp], QP_FOLDERS)
        else:
            st.warning("Please enter a keyword first.")

    if st.session_state.search_results_qp:
        st.write(f"Found **{len(st.session_state.search_results_qp)}** matching Question Paper page(s):")
        for idx, item in enumerate(st.session_state.search_results_qp):
            title_label = f"📄 {item['file']} — Page {item['page'] + 1}"
            with st.expander(title_label):
                col_preview, col_actions = st.columns([3, 2])
                with col_preview:
                    img_data = render_page_image(item['path'], item['page'])
                    st.image(img_data, caption=f"Preview of Page {item['page'] + 1}", use_container_width=True)
                with col_actions:
                    st.subheader("Actions")
                    if st.button("➕ Add to Handout Basket", key=f"add_qp_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to your basket!")
                        st.rerun()
                    st.divider()
                    with open(item['path'], "rb") as pdf_file:
                        st.download_button("📥 Download Full PDF", pdf_file, file_name=item['file'], mime="application/pdf", key=f"dl_qp_{idx}")

# --- TAB 2: SEARCH MARK SCHEMES ---
with tab2:
    st.header("Search Mark Schemes Only")
    col_input, col_reset = st.columns([4, 1])
    with col_input:
        keywords_ms = st.text_input("Enter keywords (e.g., 'evaluation')", key="input_ms")
    with col_reset:
        st.write(" ")
        if st.button("🔄 Clear Search", key="btn_clear_ms"):
            st.session_state.search_results_ms = []
            st.rerun()

    if st.button("Search Mark Schemes", type="primary", key="btn_search_ms"):
        if keywords_ms:
            with st.spinner("Scanning Mark Schemes..."):
                st.session_state.search_results_ms = search_pdfs([keywords_ms], MS_FOLDERS)
        else:
            st.warning("Please enter a keyword first.")

    if st.session_state.search_results_ms:
        st.write(f"Found **{len(st.session_state.search_results_ms)}** matching Mark Scheme page(s):")
        for idx, item in enumerate(st.session_state.search_results_ms):
            title_label = f"📄 {item['file']} — Page {item['page'] + 1}"
            with st.expander(title_label):
                col_preview, col_actions = st.columns([3, 2])
                with col_preview:
                    img_data = render_page_image(item['path'], item['page'])
                    st.image(img_data, caption=f"Preview of Page {item['page'] + 1}", use_container_width=True)
                with col_actions:
                    st.subheader("Actions")
                    if st.button("➕ Add to Handout Basket", key=f"add_ms_{idx}"):
                        st.session_state.handout_basket.append(item)
                        st.toast(f"Added Page {item['page'] + 1} to your basket!")
                        st.rerun()
                    st.divider()
                    with open(item['path'], "rb") as pdf_file:
                        st.download_button("📥 Download Full PDF", pdf_file, file_name=item['file'], mime="application/pdf", key=f"dl_ms_{idx}")

# --- TAB 3: QUICK VIEW PAPERS (UPDATED FOR INSERTS) ---
with tab3:
    st.header("Quick Download: Full Papers")
    c1, c2, c3 = st.columns(3)
    with c1:
        v_year = st.selectbox("Year", [str(y) for y in range(2029, 2018, -1)])
    with c2:
        v_month = st.selectbox("Month", ["June", "Nov"])
    with c3:
        v_paper = st.selectbox("Paper Variant", ["11", "12", "13", "21", "22", "23"])

    qp_name = get_filename_pattern(v_month, v_year, "qp", v_paper) + ".pdf"
    ms_name = get_filename_pattern(v_month, v_year, "ms", v_paper) + ".pdf"
    in_name = get_filename_pattern(v_month, v_year, "in", v_paper) + ".pdf"

    col_q, col_m, col_i = st.columns(3)
    
    # 1. Question Paper Column
    with col_q:
        path_qp = os.path.join(FOLDERS[f"{v_month} QP"], qp_name)
        if os.path.exists(path_qp):
            st.success(f"Found QP: {qp_name}")
            with open(path_qp, "rb") as f:
                st.download_button("Download Full QP", f, file_name=qp_name, key="dl_qp_tab3")
        else:
            st.error("Question Paper not found.")

    # 2. Mark Scheme Column
    with col_m:
        path_ms = os.path.join(FOLDERS[f"{v_month} MS"], ms_name)
        if os.path.exists(path_ms):
            st.success(f"Found MS: {ms_name}")
            with open(path_ms, "rb") as f:
                st.download_button("Download Full MS", f, file_name=ms_name, key="dl_ms_tab3")
        else:
            st.error("Mark Scheme not found.")

    # 3. Insert Paper Column
    with col_i:
        path_in = os.path.join(FOLDERS["Inserts"], in_name)
        if os.path.exists(path_in):
            st.success(f"Found Insert: {in_name}")
            with open(path_in, "rb") as f:
                st.download_button("Download Full Insert", f, file_name=in_name, key="dl_in_tab3")
        else:
            st.error("Insert Paper not found.")

# --- TAB 4: BASKET & EXPORT ---
with tab4:
    st.header("Worksheet Export")
    if not st.session_state.handout_basket:
        st.info("Your basket is empty. Add items from the Search tabs.")
    else:
        st.subheader(f"Items in basket: {len(st.session_state.handout_basket)}")
        
        col_btn1, col_btn2 = st.columns([1, 4])
        with col_btn1:
            if st.button("🗑️ Empty Basket", key="btn_empty_basket"):
                st.session_state.handout_basket = []
                st.rerun()

        st.write("---")
        
        st.markdown("### 📋 Selected Pages List")
        for idx, item in enumerate(st.session_state.handout_basket):
            c_label, c_del = st.columns([4, 1])
            with c_label:
                st.write(f"**{idx + 1}.** 📄 `{item['file']}` — **Page {item['page'] + 1}**")
            with c_del:
                if st.button("❌ Remove", key=f"del_item_{idx}"):
                    st.session_state.handout_basket.pop(idx)
                    st.toast("Item removed from basket!")
                    st.rerun()

        st.write("---")

        if st.button("🪄 Generate Word Handout (.docx)", type="primary", key="btn_gen_docx"):
            doc = Document()
            
            section = doc.sections[0]
            section.page_width = Inches(8.5)
            section.page_height = Inches(11.5)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.5)
            section.right_margin = Inches(0.5)
            
            add_page_number_to_header(section)

            for i, item in enumerate(st.session_state.handout_basket):
                heading = doc.add_heading(f"Source: {item['file']} (Page {item['page'] + 1})", level=2)
                heading.paragraph_format.space_before = Pt(0)
                heading.paragraph_format.space_after = Pt(4)
                
                img_data = render_page_image(item['path'], item['page'])
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_para.paragraph_format.space_before = Pt(0)
                img_para.paragraph_format.space_after = Pt(0)
                
                run = img_para.add_run()
                run.add_picture(BytesIO(img_data), width=Inches(6.2))
                
                if i < len(st.session_state.handout_basket) - 1:
                    doc.add_page_break()

            target = "8021_Custom_Handout.docx"
            doc.save(target)
            with open(target, "rb") as f:
                st.download_button("📥 Click to Download Document", f, file_name=target, key="dl_docx")

# --- TAB 5: ADMIN CONTROL ---
with tab5:
    st.header("Admin Control Center")
    pwd = st.text_input("Enter Admin Password to access controls", type="password", key="admin_pwd")
    if pwd == ADMIN_PASSWORD:
        st.success("Welcome, Admin.")
        st.markdown("""
        ### 📂 Library Management
        To add or remove papers permanently, please use the Google Drive portal. 
        Changes will reflect here after clicking **Sync** in the sidebar.
        """)
        st.link_button("Go to Google Drive Library", f"https://drive.google.com/drive/folders/{GD_FOLDER_ID}")
        st.divider()
        st.subheader("📊 Live System Status")
        for label, folder in FOLDERS.items():
            if os.path.exists(folder):
                file_count = len([f for f in os.listdir(folder) if f.endswith('.pdf')])
                st.write(f"✅ **{label}:** {file_count} files synced")
            else:
                st.error(f"❌ **{label}:** Folder missing!")
    elif pwd:
        st.error("Incorrect Password. Access Denied.")

# --- FOOTER ---
st.markdown("---")
st.markdown(
    """
    <div style="text-align: center; width: 100%;">
        <p style="font-size: 20px; font-weight: bold; margin-bottom: 5px;">
            ✨ PTES 8021 General Paper Resource Portal ✨
        </p>
        <p style="font-size: 16px; font-weight: bold; letter-spacing: 0.5px;">
            <span style="color: #FF0000;">🔴 Academic Excellence</span> | 
            <span style="color: #FFD700;">🟡 Future Readiness</span> | 
            <span style="color: #0070FF;">🔵 Digital & Integrity</span> | 
            <span style="color: #28A745;">🟢 Holistic & Growth</span>
        </p>
        <p style="color: #057047; font-size: 14px; margin-top: 10px;">
            Creator: Miss Hajah Nurul Haziqah HN (PTES CS Tutor)
        </p>
    </div>
    """,
    unsafe_allow_html=True
)
